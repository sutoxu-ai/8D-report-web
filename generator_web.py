#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
8D 报告智能生成助手 - 客户端
最终优化版本
"""

import streamlit as st
from io import BytesIO
from datetime import datetime, timedelta
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
import openai
from supabase import create_client
import logging

# ==================== 缓存配置 ====================
@st.cache_data(ttl=60)
def get_cached_license(user_id):
    """缓存用户许可证信息，60 秒 TTL"""
    if not supabase:
        return None
    try:
        r = supabase.table("licenses").select("*").eq("user_id", user_id).execute()
        if r.data:
            return r.data[0]
        return create_free_license(user_id)
    except Exception:
        return None

def clear_license_cache(user_id):
    """清除特定用户的缓存"""
    get_cached_license.clear()

# ==================== 页面配置 ====================
st.set_page_config(
    page_title="8D 报告 - 智能生成助手", 
    page_icon="📊", 
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==================== 隐藏 Streamlit 默认 UI 元素 ====================
hide_streamlit_style = """
<style>
    /* 只隐藏右上角菜单 */
    #MainMenu {visibility: hidden !important; display: none !important;}
    
    /* 隐藏 footer 水印 */
    footer {visibility: hidden !important; display: none !important;}
    
    /* 隐藏 Pages 导航菜单列表 */
    [data-testid="stSidebarNav"] > ul {display: none !important;}
    
    /* 调整主内容区域 */
    .main .block-container {
        padding-top: 0.5rem !important;
    }
    
    /* ========== 缩小侧边栏间距 ========== */
    [data-testid="stSidebar"] .block-container {
        padding-top: 1rem !important;
        padding-bottom: 1rem !important;
    }
    
    [data-testid="stSidebar"] h3 {
        margin-top: 0.5rem !important;
        margin-bottom: 0.3rem !important;
    }
    
    [data-testid="stSidebar"] p {
        margin-top: 0.2rem !important;
        margin-bottom: 0.2rem !important;
    }
    
    [data-testid="stSidebar"] .stButton {
        margin-top: 0.2rem !important;
        margin-bottom: 0.2rem !important;
    }
    
    [data-testid="stSidebar"] .stTextInput {
        margin-top: 0.2rem !important;
        margin-bottom: 0.2rem !important;
    }
    
    [data-testid="stSidebar"] .streamlit-expanderHeader {
        padding-top: 0.3rem !important;
        padding-bottom: 0.3rem !important;
    }
    
    [data-testid="stSidebar"] .stCaption {
        margin-top: 0.1rem !important;
        margin-bottom: 0.1rem !important;
    }
    
    [data-testid="stSidebar"] .stMarkdown {
        margin-bottom: 0.3rem !important;
    }
    
    [data-testid="stSidebar"] div[data-testid="stVerticalBlock"] > div {
        gap: 0.15rem !important;
    }
    
    [data-testid="stSidebar"] hr {
        display: none !important;
    }
    
    /* ========== 手机端字体缩小 ========== */
    @media screen and (max-width: 768px) {
        h1 { font-size: 1.3rem !important; }
        h2 { font-size: 1.1rem !important; }
        h3 { font-size: 1rem !important; }
        h4 { font-size: 0.95rem !important; }
        body { font-size: 0.85rem !important; }
        input, textarea { font-size: 0.85rem !important; }
        button { font-size: 0.9rem !important; }
        label, .stMarkdown p, .stMarkdown span { font-size: 0.85rem !important; }
        /* 手机端缩小文本输入框高度 */
        textarea {
            height: 80px !important;
            min-height: 80px !important;
        }
    }
</style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# ==================== JavaScript 隐藏右上角按钮 ====================
hide_buttons_script = """
<script>
// Hide top-right toolbar buttons, keep sidebar toggle
function hideTopRightButtons() {
    // Hide links containing fork, github, star
    document.querySelectorAll('a').forEach(el => {
        const href = (el.href || '').toLowerCase();
        const text = (el.textContent || '').toLowerCase();
        if (href.includes('fork') || href.includes('github') || 
            text.includes('fork') || text.includes('star')) {
            el.style.display = 'none';
        }
    });
}

if (document.readyState === 'loading') {
    document.addEventListener('DOMContentLoaded', () => setTimeout(hideTopRightButtons, 500));
} else {
    setTimeout(hideTopRightButtons, 500);
}
</script>
"""
st.markdown(hide_buttons_script, unsafe_allow_html=True)

# ==================== 多语言文本 ====================
TEXT = {
    "zh": {
        "lang_label": "语言", "lang_zh": "中文", "lang_en": "English",
        "system_status": "系统状态", "pro_version": "✅ 正式版", "trial_version": "⚠️ 试用版",
        "license_valid_until": "📅 有效期至 {exp}", "trial_used": "📊 已使用 {used} 次 / 共 {total} 次",
        "trial_exhausted": "❌ 试用次数已用完", "activate_title": "🔑 授权 / 续费",
        "activate_code_hint": "激活码", "activate_btn": "立即激活",
        "activate_success": "✅ 激活成功，有效期一年", "activate_fail": "❌ 激活码无效",
        "invalid_activate_code": "请输入有效的激活码",
        "license_expired": "❌ 授权已过期",
        "login_required": "🔒 请先登录", "logout": "退出登录",
        "login_header": "👤 用户登录",
        "username_label": "邮箱或手机号",
        "username_placeholder": "例：zhangsan@163.com 或 13812345678",
        "login_register_btn": "🔓 登录 / 注册",
        "enter_username_error": "请输入邮箱或手机号",
        "invalid_email": "❌ 邮箱格式不正确，示例：zhangsan@163.com",
        "invalid_phone": "❌ 手机号格式不正确，请输入11位大陆手机号（1开头，第二位3-9）",
        "invalid_contact": "❌ 请输入有效的邮箱或11位大陆手机号",
        "expander_activate_code": "🔑 输入激活码",
        "enter_activate_code_placeholder": "输入激活码",
        "trial_remaining": "📊 **试用版** | 剩余 {n} 次",
        "no_trial_hint": "💡 试用次数用完了？",
        "valid_until": "⏰ 有效期至: {date}",
        "valid_until_date": "📅 有效期至: {date}",
        "permanent_valid": "♾️ 永久有效",
        "account_manager": "🔐 账户管理 / Account",
        "contact_service": "📱 联系客服 / Contact",
        "new_user_hint": "👋 注册后可购买试用券",
        "main_title": "📊 8D 报告智能生成助手",
        "progress_phases": [
            {"icon": "📝", "text": "正在整理您的输入信息...", "sub": "产品：{product}"},
            {"icon": "🤔", "text": "正在理解问题背景...", "sub": "运用 5W2H 方法分析"},
            {"icon": "📋", "text": "正在生成 D1 团队组建...", "sub": "确定责任人及时间节点"},
            {"icon": "📋", "text": "正在生成 D2 问题描述...", "sub": "详细记录不良现象"},
            {"icon": "🛡️", "text": "正在生成 D3 临时措施...", "sub": "遏制问题扩散"},
            {"icon": "🔬", "text": "正在分析根本原因 (4M1E)...", "sub": "人、机、料、法、环逐一排查"},
            {"icon": "🔍", "text": "正在进行 5-Why 追问...", "sub": "追溯至根本原因"},
            {"icon": "💡", "text": "正在制定 D5 永久措施...", "sub": "根本性解决方案"},
            {"icon": "✅", "text": "正在生成 D6 实施计划...", "sub": "验证措施有效性"},
            {"icon": "📊", "text": "正在生成 D7 预防措施...", "sub": "防止问题复发"},
            {"icon": "🏆", "text": "正在生成 D8 总结表彰...", "sub": "固化经验，分享成果"},
            {"icon": "✨", "text": "正在优化报告格式...", "sub": "确保专业美观"},
        ], "input_header": "📝 输入基本信息",
        "product_name": "产品型号 / 名称", "customer": "客户名称",
        "problem_desc": "不良现象描述",
        "problem_placeholder": "请使用 5W2H 方法描述问题",
        "occur_date": "发现日期", "defect_qty": "不良数量", "severity": "严重程度",
        "severity_low": "低", "severity_medium": "中", "severity_high": "高", "severity_critical": "危急",
        "industry_std": "适用标准", "team_members": "团队成员（可选）",
        "team_placeholder": "例：张明 (组长), 李华 (工程)",
        "generate_btn": "🚀 生成 8D 报告", 
        "generating": "8D 报告智能生成中，请稍候...",
        "preview_header": "📄 报告预览", "download_btn": "📥 导出 Word 报告",
        "export_disabled": "🔒 激活正式版后可导出 Word",
        "no_desc": "❌ 请输入不良现象描述",
        "trial_exhausted_error": "❌ 试用次数已用完", "api_error": "❌ 服务异常",
        "success": "✅ 报告生成完成！", "report_complete": "报告生成完成！",
        "beautifying": "正在美化格式...", "word_title": "8D 问题纠正与预防措施报告",
        "system_error": "❌ 系统错误，请稍后重试",
        "history_header": "📋 生成历史",
        "load_report": "加载此报告",
        "delete_report": "删除",
        "no_history": "暂无历史记录",
        "edit_mode": "✏️ 编辑模式",
        "save_edit": "💾 保存修改",
        "edit_placeholder": "在此编辑报告内容...",
        "history_loaded": "已从历史记录加载",
    },
    "en": {
        "lang_label": "Language", "lang_zh": "中文", "lang_en": "English",
        "system_status": "System Status", "pro_version": "✅ Pro Version", "trial_version": "⚠️ Trial Version",
        "license_valid_until": "📅 Valid until {exp}", "trial_used": "📊 Used {used} / {total}",
        "trial_exhausted": "❌ Trial exhausted", "activate_title": "🔑 License / Renew",
        "activate_code_hint": "Activation Code", "activate_btn": "Activate",
        "activate_success": "✅ Activated successfully", "activate_fail": "❌ Invalid code",
        "invalid_activate_code": "Please enter a valid activation code",
        "license_expired": "❌ License expired",
        "login_required": "🔒 Please login", "logout": "Logout",
        "login_header": "👤 User Login",
        "username_label": "Email or Phone",
        "username_placeholder": "e.g., name@example.com or 13812345678",
        "login_register_btn": "🔓 Login / Register",
        "enter_username_error": "Please enter email or phone number",
        "invalid_email": "❌ Invalid email format, e.g., name@example.com",
        "invalid_phone": "❌ Invalid phone number, enter 11-digit mainland China number",
        "invalid_contact": "❌ Please enter a valid email or 11-digit phone number",
        "expander_activate_code": "🔑 Enter Activation Code",
        "enter_activate_code_placeholder": "Enter activation code",
        "trial_remaining": "📊 **Trial** | {n} remaining",
        "no_trial_hint": "💡 Run out of trials?",
        "valid_until": "⏰ Valid until: {date}",
        "valid_until_date": "📅 Valid until: {date}",
        "permanent_valid": "♾️ Permanent",
        "account_manager": "🔐 Account Manager",
        "contact_service": "📱 Contact Service",
        "new_user_hint": "👋 Register to get started, then purchase trial",
        "main_title": "📊 8D Report Generator",
        "progress_phases": [
            {"icon": "📝", "text": "Organizing your input...", "sub": "Product: {product}"},
            {"icon": "🤔", "text": "Analyzing context...", "sub": "Using 5W2H method"},
            {"icon": "📋", "text": "Generating D1 Team...", "sub": "Defining responsibilities"},
            {"icon": "📋", "text": "Generating D2 Description...", "sub": "Recording defect details"},
            {"icon": "🛡️", "text": "Generating D3 Containment...", "sub": "Preventing spread"},
            {"icon": "🔬", "text": "Analyzing root cause (4M1E)...", "sub": "Checking all factors"},
            {"icon": "🔍", "text": "Performing 5-Why analysis...", "sub": "Finding root cause"},
            {"icon": "💡", "text": "Developing D5 Actions...", "sub": "Long-term solutions"},
            {"icon": "✅", "text": "Generating D6 Implementation...", "sub": "Verifying effectiveness"},
            {"icon": "📊", "text": "Generating D7 Prevention...", "sub": "Preventing recurrence"},
            {"icon": "🏆", "text": "Generating D8 Closure...", "sub": "Documenting lessons"},
            {"icon": "✨", "text": "Formatting report...", "sub": "Professional output"},
        ], "input_header": "📝 Input Information",
        "product_name": "Product Name / Model", "customer": "Customer Name",
        "problem_desc": "Problem Description",
        "problem_placeholder": "Please use 5W2H method",
        "occur_date": "Occurrence Date", "defect_qty": "Defect Quantity", "severity": "Severity",
        "severity_low": "Low", "severity_medium": "Medium", "severity_high": "High", "severity_critical": "Critical",
        "industry_std": "Standard", "team_members": "Team Members (Optional)",
        "team_placeholder": "e.g., Zhang(Leader), Li(Eng)",
        "generate_btn": "🚀 Generate 8D Report",
        "generating": "Generating report, please wait...",
        "preview_header": "📄 Report Preview", "download_btn": "📥 Export Word",
        "export_disabled": "🔒 Activate to export",
        "no_desc": "❌ Please enter description",
        "trial_exhausted_error": "❌ Trial exhausted", "api_error": "❌ Service error",
        "success": "✅ Report generated!", "report_complete": "Report generated!",
        "beautifying": "Formatting...", "word_title": "8D Corrective Action Report",
        "system_error": "❌ System error, please try again later",
        "history_header": "📋 Generation History",
        "load_report": "Load this report",
        "delete_report": "Delete",
        "no_history": "No history yet",
        "edit_mode": "✏️ Edit Mode",
        "save_edit": "💾 Save Changes",
        "edit_placeholder": "Edit report content here...",
        "history_loaded": "Loaded from history",
    }
}

# ==================== 系统提示词 ====================
SYSTEM_PROMPT = {
    "zh": (
        "你是一位拥有 20 年经验的汽车电子行业高级质量工程师，精通 IATF 16949 标准和 8D 问题解决方法。"
        "请根据用户输入撰写专业、逻辑严密的 8D 报告。\n\n"
        "【8D 报告结构要求】\n"
        "报告必须严格按照以下 8 个步骤的顺序输出，不可颠倒：\n"
        "D1：建立团队（成立问题解决小组，列出成员及职责）\n"
        "D2：问题描述（使用 5W2H 方法描述问题：What、Why、Who、When、Where、How、How many）\n"
        "D3：制定临时控制措施（ICA，围堵措施，防止问题扩大）\n"
        "D4：根本原因分析（见下方详细要求）\n"
        "D5：制定永久纠正措施（PCA，针对根本原因的根本解决方案）\n"
        "D6：贯彻永久纠正措施（实施计划、验证有效性）\n"
        "D7：预防措施（防止类似问题在其他产品/流程中复发）\n"
        "D8：表彰小组（总结、表彰团队成员贡献）\n\n"
        "【D4 根本原因分析要求】\n"
        "根本原因分析必须包含两部分：产生原因 和 流出原因。\n\n"
        "一、产生原因分析（为什么会产生缺陷）：\n"
        "使用 4M1E 分析法（人、机、料、法、环）逐项确认，使用确定句而非疑问句：\n"
        "✅ 正常项：明确说明\"经检查，XX 符合标准，排除为根本原因\"\n"
        "❌ 异常项：明确说明\"经检查，XX 存在问题：[具体问题]\"\n"
        "❌ 不要使用\"是否\"、\"有没有\"等疑问句\n\n"
        "从异常项开始，使用 5-Why 分析法：\n"
        "连续追问\"为什么\"，至少追问 3-5 层，直到找到根本原因\n"
        "每层回答要具体，不能笼统\n\n"
        "二、流出原因分析（为什么缺陷没有被发现，流向了客户）：\n"
        "分析检验/拦截环节为什么会失效，同样使用 5-Why 分析法：\n"
        "Why1：为什么该缺陷在 XX 检验环节没有被发现？\n"
        "Why2：为什么检验标准/方法/频次存在漏洞？\n"
        "Why3：为什么检验人员没有执行到位？\n"
        "Why4：为什么检验流程设计不完善？\n"
        "Why5：为什么管理层没有重视检验环节？\n\n"
        "输出格式示例（注意换行）：\n"
        "【D4 根本原因分析】\n\n"
        "=== 产生原因分析 ===\n\n"
        "【4M1E 分析】\n\n"
        "人：经检查，操作员持证上岗 → 排除\n\n"
        "机：经检查，设备参数偏移 0.05mm → 异常项 ⚠️\n\n"
        "料：经检查，原材料合格 → 排除\n\n"
        "法：经检查，作业指导书过期 → 异常项 ⚠️\n\n"
        "环：经检查，环境符合要求 → 排除\n\n"
        "【5-Why 分析（产生原因）】\n\n"
        "Why1：为什么设备参数偏移？→ 传感器校准过期\n\n"
        "Why2：为什么校准过期？→ 年度校准计划未执行\n\n"
        "Why3：为什么计划未执行？→ 维护人员不足 ← 根本原因\n\n"
        "=== 流出原因分析 ===\n\n"
        "【检验环节失效分析】\n\n"
        "检验点：出货检验 OQC\n\n"
        "Why1：为什么偏移参数的产品流出了？→ OQC 检验标准未包含该参数\n\n"
        "Why2：为什么标准未包含？→ 控制计划未更新该参数\n\n"
        "Why3：为什么控制计划未更新？→ 工程变更流程缺失 ← 根本原因\n\n"
        "【其他要求】\n"
        "语气专业客观\n"
        "措施使用 [责任人 | 时间 | 状态] 格式\n"
        "不使用 Markdown 标记\n"
        "直接输出 D1-D8 报告正文"
    ),
    
    "en": (
        "You are a Senior Quality Engineer with 20 years experience in automotive electronics, "
        "proficient in IATF 16949 and 8D methodology. Please write a professional 8D report based on user input.\n\n"
        "【8D Report Structure Requirements】\n"
        "The report MUST follow these 8 steps in strict order, do not swap them:\n"
        "D1: Establish Team (form the problem solving team, list members and roles)\n"
        "D2: Describe the Problem (use 5W2H: What, Why, Who, When, Where, How, How many)\n"
        "D3: Develop Interim Containment Plan (ICA, containment actions to prevent spread)\n"
        "D4: Root Cause Analysis (see detailed requirements below)\n"
        "D5: Develop Permanent Corrective Actions (PCA, solutions addressing root cause)\n"
        "D6: Implement and Validate Corrective Actions (implementation plan, verify effectiveness)\n"
        "D7: Preventive Measures (prevent recurrence in similar products/processes)\n"
        "D8: Recognize Team and Individual Contributions (conclude and recognize team)\n\n"
        "【D4 Root Cause Analysis Requirements】\n"
        "Root cause analysis MUST include two parts: Occurrence Cause and Escape Cause.\n\n"
        "Part 1 - Occurrence Cause (Why did the defect occur?):\n"
        "Use 4M1E analysis (Man, Machine, Material, Method, Environment) with declarative sentences:\n"
        "✅ Normal: 'Verified, XX meets standard, excluded as root cause'\n"
        "❌ Abnormal: 'Verified, XX has issue: [specific problem]'\n"
        "Then use 5-Why analysis from abnormal items, minimum 3-5 levels.\n\n"
        "Part 2 - Escape Cause (Why did the defect escape to customer?):\n"
        "Analyze why inspection/containment failed, use 5-Why:\n"
        "Why1: Why wasn't this defect caught at XX inspection?\n"
        "Why2: Why is there a gap in inspection standard/method/frequency?\n"
        "Why3: Why wasn't the inspector performing correctly?\n"
        "Why4: Why is the inspection process flawed?\n"
        "Why5: Why didn't management prioritize this?\n\n"
        "Format example:\n"
        "【D4 Root Cause Analysis】\n\n"
        "=== Occurrence Cause ===\n\n"
        "【4M1E Analysis】\n"
        "Man: Verified... Excluded\n\n"
        "【5-Why Analysis (Occurrence)】\n"
        "Why1: ...\n\n"
        "=== Escape Cause ===\n\n"
        "【Inspection Failure Analysis】\n"
        "Why1: ...\n\n"
        "【Other Requirements】\n"
        "Professional tone\n"
        "Use [Owner|Date|Status] format for actions\n"
        "No Markdown\n"
        "Output D1-D8 directly"
    )
}
# ==================== 初始化配置 ====================
try:
    API_KEY = st.secrets["DEEPSEEK_API_KEY"]
    BASE_URL = st.secrets["DEEPSEEK_BASE_URL"]
except Exception:
    API_KEY = ""
    BASE_URL = "https://api.deepseek.com"

try:
    supabase = create_client(st.secrets["SUPABASE_URL"], st.secrets["SUPABASE_KEY"])
except Exception:
    supabase = None

# ==================== 核心功能函数 ====================
def get_user_license(user_id):
    return get_cached_license(user_id)

def create_free_license(user_id):
    if not supabase:
        return None
    try:
        r = supabase.table("licenses").insert({
            "user_id": user_id,
            "plan_type": "free",
            "trial_used": 0,
            "trial_limit": 0
        }).execute()
        return r.data[0] if r.data else None
    except Exception:
        return None

def can_generate_report(user_id):
    lic = get_user_license(user_id)
    if not lic:
        return False
    if lic['plan_type'] in ['free', 'trial']:
        return lic['trial_used'] < lic['trial_limit']
    if lic['plan_type'] in ['pro', 'enterprise']:
        if lic.get('license_expire'):
            try:
                return datetime.now() < datetime.fromisoformat(lic['license_expire'])
            except Exception:
                return True
        return True
    return False

def inc_trial_used(user_id):
    if not supabase:
        return
    try:
        lic = get_cached_license(user_id)
        if lic:
            new_count = lic.get('trial_used', 0) + 1
            supabase.table("licenses").update({"trial_used": new_count}).eq("user_id", user_id).execute()
            supabase.table("usage_logs").insert({
                "user_id": user_id,
                "action": "generate_report",
                "created_at": datetime.now().isoformat()
            }).execute()
            clear_license_cache(user_id)
    except Exception as e:
        logging.error(f"更新试用次数失败：{e}")

def activate_license_code(user_id, code):
    if not supabase:
        return False, "系统错误"
    try:
        r = supabase.table("activation_codes").select("*").eq("code", code.strip().upper()).execute()
        if not r.data:
            return False, "无效的激活码"
        ac = r.data[0]
        if ac.get('is_used'):
            return False, "激活码已被使用"
        if ac.get('expire_date'):
            if datetime.now().date() > datetime.fromisoformat(ac['expire_date']).date():
                return False, "激活码已过期"
        duration = ac.get('duration_days', 365)
        exp_date = (datetime.now() + timedelta(days=duration)).isoformat()
        supabase.table("licenses").upsert({
            "user_id": user_id,
            "plan_type": ac.get('plan_type', 'pro'),
            "license_expire": exp_date
        }, on_conflict="user_id").execute()
        supabase.table("activation_codes").update({
            "is_used": True,
            "used_by": user_id,
            "used_at": datetime.now().isoformat()
        }).eq("code", code.strip().upper()).execute()
        clear_license_cache(user_id)
        formatted_date = exp_date[:10] if len(exp_date) >= 10 else exp_date
        return True, f"激活成功！有效期至 {formatted_date}"
    except Exception as e:
        logging.error(f"激活失败：{e}")
        return False, f"激活失败：{str(e)}"

def activate_trial_code(user_id, code):
    """激活试用码：0.99元获得2次试用"""
    if not supabase:
        return False, "系统错误"
    try:
        r = supabase.table("trial_codes").select("*").eq("code", code.strip().upper()).execute()
        if not r.data:
            return False, "无效的试用码"
        tc = r.data[0]
        if tc.get('is_used'):
            return False, "试用码已被使用"
        # 标记试用码已用
        supabase.table("trial_codes").update({
            "is_used": True,
            "used_by": user_id,
            "used_at": datetime.now().isoformat()
        }).eq("code", code.strip().upper()).execute()
        # 给用户2次试用机会
        supabase.table("licenses").update({
            "trial_limit": 2,
            "trial_used": 0
        }).eq("user_id", user_id).execute()
        clear_license_cache(user_id)
        return True, "✅ 试用码激活成功！获得 2 次试用机会"
    except Exception as e:
        logging.error(f"试用码激活失败：{e}")
        return False, f"激活失败：{str(e)}"

def clean_format(text):
    if not text:
        return ""
    text = text.replace("**", "").replace("#", "")
    for i in range(1, 9):
        text = re.sub(rf'(D{i}[:：])\s*\n+\s*', rf'\1 ', text)
    text = re.sub(r'([人机料法环]：)', r'\n\1', text)
    text = re.sub(r'(→ 排除|→ 异常项[^，]*？)', r'\1\n', text)
    text = re.sub(r'(Why\d+：)', r'\n\1', text)
    text = re.sub(r'(→ [^\n]+)(?=Why\d+：|$)', r'\1\n', text)
    text = re.sub(r'(为什么\d+：)', r'\n\1', text)
    text = re.sub(r'(→ [^\n]+)(?=为什么\d+：|$)', r'\1\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def export_to_word(content, product_name, lang):
    doc = Document()
    if lang == "zh":
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    else:
        doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10.5)
    title = doc.add_heading(TEXT[lang]["word_title"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info = doc.add_paragraph()
    info.add_run(f"Product: {product_name}").bold = True
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    sections = re.split(r'\n(?=D[1-8][:：])', content.replace("**", "").replace("#", ""))
    for i, sec in enumerate(sections):
        if not sec.strip():
            continue
        lines = sec.strip().split('\n', 1)
        p_title = doc.add_paragraph()
        runner = p_title.add_run(lines[0].strip())
        runner.bold = True
        runner.font.size = Pt(14)
        runner.font.color.rgb = RGBColor(30, 58, 138)
        if len(lines) > 1 and lines[1].strip():
            doc.add_paragraph(lines[1].strip())
        if i < len(sections) - 1:
            p_line = doc.add_paragraph()
            p_line.paragraph_format.space_before = Pt(12)
            p_line.paragraph_format.space_after = Pt(12)
            p = p_line._element
            pPr = p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '8')
            pBdr.append(bottom)
            pPr.append(pBdr)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ==================== 历史记录功能 ====================
def save_report_history(user_id, product_name, customer, problem_desc, report_content, lang):
    """保存报告到历史记录"""
    if not supabase or not user_id:
        return
    try:
        supabase.table("reports").insert({
            "user_id": user_id,
            "product_name": product_name or "",
            "customer": customer or "",
            "problem_desc": (problem_desc or "")[:500],
            "report_content": report_content,
            "lang": lang,
            "created_at": datetime.now().isoformat()
        }).execute()
    except Exception as e:
        logging.warning(f"保存历史记录失败：{e}")

def load_report_history(user_id, limit=10):
    """加载用户历史记录"""
    if not supabase or not user_id:
        return []
    try:
        r = supabase.table("reports").select("*").eq("user_id", user_id).order("created_at", desc=True).limit(limit).execute()
        return r.data or []
    except Exception as e:
        logging.warning(f"加载历史记录失败：{e}")
        return []

def delete_report_history(report_id):
    """删除单条历史记录"""
    if not supabase:
        return False
    try:
        supabase.table("reports").delete().eq("id", report_id).execute()
        return True
    except Exception as e:
        logging.warning(f"删除历史记录失败：{e}")
        return False

# ==================== 会话状态初始化 ====================
if "lang" not in st.session_state:
    st.session_state.lang = "zh"
if "current_result" not in st.session_state:
    st.session_state.current_result = ""
if "user_id" not in st.session_state:
    st.session_state.user_id = None
if "registration_attempted" not in st.session_state:
    st.session_state.registration_attempted = False

T = TEXT[st.session_state.lang]

# ==================== 侧边栏（登录和用户管理） ====================
def render_sidebar():
    """渲染侧边栏 - 登录、用户信息、语言切换、激活码等"""
    T = TEXT[st.session_state.lang]
    
    with st.sidebar:
        # ==================== 语言切换 ====================
        st.markdown("### 🌐 语言 / Language")
        lang_option = st.selectbox(
            "选择语言 / Select Language",
            ["中文", "English"],
            index=0 if st.session_state.lang == "zh" else 1,
            key="sidebar_lang_select",
            label_visibility="collapsed"
        )
        new_lang = "zh" if lang_option == "中文" else "en"
        if new_lang != st.session_state.lang:
            st.session_state.lang = new_lang
            st.rerun()
        
        st.markdown(f"### {T['account_manager']}")
                
        # ==================== 登录 / 注册区域 ====================
        user_id = st.session_state.get("user_id")
        
        if not user_id:
            st.info(T["new_user_hint"])
            
            user_input = st.text_input(
                T["username_label"],
                key="sidebar_user_input",
                placeholder=T["username_placeholder"]
            )
            
            # ========== 格式校验函数 ==========
            def validate_contact(contact):
                """校验邮箱或手机号格式"""
                email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
                phone_pattern = r'^1[3-9]\d{9}$'
                
                if re.match(email_pattern, contact):
                    return True, "email"
                elif re.match(phone_pattern, contact):
                    return True, "phone"
                else:
                    return False, None
            
            if st.button(T["login_register_btn"], use_container_width=True, key="sidebar_login_btn"):
                if not user_input:
                    st.error(T["enter_username_error"])
                    st.stop()

                # ========== 方案1：会话级注册限制 ==========
                if st.session_state.registration_attempted:
                    st.error("⚠️ 当前会话已注册过，请勿重复操作" if st.session_state.lang == "zh" else "⚠️ Already registered in this session")
                    st.stop()

                # 检查是否为已注册的老用户
                existing_user = False
                if supabase:
                    try:
                        existing = supabase.table("licenses").select("user_id").eq("user_id", user_input).execute()
                        existing_user = existing.data is not None and len(existing.data) > 0
                    except Exception:
                        st.error(T["system_error"])
                        st.stop()

                # 校验格式
                is_valid, contact_type = validate_contact(user_input)

                # 老用户不受格式限制，直接放行
                if not existing_user and not is_valid:
                    if "@" in user_input:
                        st.error(T["invalid_email"])
                    elif user_input.startswith("1") and len(user_input) == 11:
                        st.error(T["invalid_phone"])
                    else:
                        st.error(T["invalid_contact"])
                    st.stop()

                # ========== 方案2：相似账号检测（仅新用户） ==========
                if not existing_user and supabase and user_input.isdigit() and len(user_input) >= 10:
                    try:
                        prefix_len = len(user_input) - 2
                        prefix = user_input[:prefix_len]
                        similar = supabase.table("licenses").select("user_id").like("user_id", prefix + "%").limit(5).execute()
                        if similar.data and len(similar.data) > 0:
                            st.error(
                                "⚠️ 检测到可疑注册行为，已被拒绝。请联系客服。"
                                if st.session_state.lang == "zh" else
                                "⚠️ Suspicious registration detected. Please contact support."
                            )
                            st.stop()
                    except Exception as e:
                        logging.warning(f"相似账号检测失败：{e}")

                # 新用户注册（格式验证通过 且 无历史记录）
                if not existing_user and supabase:
                    try:
                        supabase.table("licenses").insert({
                            "user_id": user_input,
                            "plan_type": "free",
                            "trial_used": 0,
                            "trial_limit": 0
                        }).execute()
                        st.session_state.registration_attempted = True  # 标记会话已注册
                    except Exception:
                        pass

                st.session_state.user_id = user_input
                st.rerun()
        
        # ==================== 已登录用户区域 ====================
        else:
            lic = get_user_license(user_id)
            
            st.markdown(f"**👤 {user_id[:30]}**")
            
            if lic:
                if lic.get('plan_type') == 'free':
                    remaining = (lic.get('trial_limit') or 0) - (lic.get('trial_used') or 0)
                    if remaining > 0:
                        st.info(T["trial_remaining"].format(n=remaining))
                    else:
                        st.error("❌ 试用次数已用完")
                        # 试用购买引导
                        st.markdown("---")
                        st.markdown("### 💰 购买试用券")
                        st.caption("¥0.99 = 2 次试用")
                        st.info("""
**购买步骤：**
1. 微信转账 **¥0.99** 到客服
2. 截图发给微信 **907749064**
3. 客服发送**试用码**
4. 输入试用码获得 2 次试用
                        """)
                        # 试用码输入
                        trial_code_input = st.text_input(
                            "输入试用码",
                            type="password",
                            key="sidebar_trial_code",
                            placeholder="例：8DT1-XXXX-XXXX-X"
                        )
                        if st.button("激活试用码", key="sidebar_trial_btn", use_container_width=True):
                            if trial_code_input and len(trial_code_input) >= 6:
                                success, msg = activate_trial_code(user_id, trial_code_input)
                                if success:
                                    st.success(msg)
                                    st.rerun()
                                else:
                                    st.error(msg)
                            else:
                                st.error("请输入有效的试用码")
                else:
                    st.success(T["pro_version"])
                    if lic.get('license_expire'):
                        try:
                            exp_date = datetime.fromisoformat(lic['license_expire']).strftime('%Y-%m-%d')
                            st.caption(T["valid_until_date"].format(date=exp_date))
                        except:
                            pass
                    else:
                        st.caption(T["permanent_valid"])
            
                        
            with st.expander(T["expander_activate_code"], expanded=False):
                activate_code = st.text_input(
                    T["activate_code_hint"],
                    type="password",
                    key="sidebar_act_code",
                    placeholder=T["enter_activate_code_placeholder"]
                )
                if st.button(T["activate_btn"], key="sidebar_act_btn", use_container_width=True):
                    if activate_code and len(activate_code) >= 6:
                        success, msg = activate_license_code(user_id, activate_code)
                        if success:
                            st.success(msg)
                            st.rerun()
                        else:
                            st.error(msg)
                    else:
                        st.error(T["invalid_activate_code"])
            
            if st.button(T["logout"], key="sidebar_logout_btn", use_container_width=True):
                st.session_state.user_id = None
                st.session_state.current_result = ""
                get_cached_license.clear()
                st.rerun()
        
        # ==================== 历史记录 ====================
        if user_id:
            st.markdown(f"**{T['history_header']}**")
            history = load_report_history(user_id)
            if not history:
                st.caption(T["no_history"])
            else:
                for report in history[:5]:
                    with st.expander(f"{(report.get('product_name') or 'N/A')[:20]} | {report['created_at'][:10]}"):
                        st.caption((report.get('problem_desc') or '')[:80])
                        col_load, col_del = st.columns([3, 1])
                        with col_load:
                            if st.button(T["load_report"], key=f"load_{report['id']}", use_container_width=True):
                                st.session_state.current_result = report['report_content']
                                st.success(T["history_loaded"])
                                st.rerun()
                        with col_del:
                            if st.button("🗑️", key=f"del_{report['id']}", use_container_width=True):
                                if delete_report_history(report['id']):
                                    st.rerun()

        st.markdown("---")

        # ==================== 底部信息 ====================
        
        st.markdown(f"**{T['contact_service']}**")
        try:
            st.image("wechat_qrcode.jpg", width=180)
        except:
            st.info("微信二维码：907749064")
        st.caption("淘宝店铺: 效率工坊铺")
        st.caption("微信号Wechat: 907749064")
        st.caption("Email: 907749064@qq.com")
        st.markdown("---")
 
# ==================== 主页面 ====================
render_sidebar()

st.title(T["main_title"])

col_input, col_preview = st.columns([1, 1.2])

with col_input:
    st.header(T["input_header"])
    product_name = st.text_input(T["product_name"], placeholder="e.g., PCB-A123" if st.session_state.lang == "en" else "例：PCB-A123")
    customer = st.text_input(T["customer"], placeholder="e.g., BYD" if st.session_state.lang == "en" else "例：比亚迪汽车")
    problem_desc = st.text_area(T["problem_desc"], height=150, placeholder=T["problem_placeholder"])
    
    col1, col2, col3 = st.columns(3)
    with col1:
        occur_date = st.date_input(T["occur_date"], datetime.now())
    with col2:
        defect_qty = st.number_input(T["defect_qty"], min_value=1, value=1)
    with col3:
        severity = st.selectbox(
            T["severity"], 
            [T["severity_low"], T["severity_medium"], T["severity_high"], T["severity_critical"]]
        )
    
    col4, col5 = st.columns(2)
    with col4:
        industry_std = st.selectbox(
            T["industry_std"], 
            ["ISO 9001", "IATF 16949", "ISO 13485", "AS9100"], 
            index=1
        )
    with col5:
        team_members = st.text_input(T["team_members"], placeholder=T["team_placeholder"])
    
        
    if st.button(T["generate_btn"], type="primary", use_container_width=True):
        if not st.session_state.get("user_id"):
            st.error(T["login_required"])
            st.stop()

        user_id = st.session_state.user_id
        if not can_generate_report(user_id):
            lic = get_user_license(user_id)
            if lic and lic['plan_type'] == 'free':
                st.error(T["trial_exhausted_error"])
            else:
                st.error(T["license_expired"])
            st.stop()

        if not problem_desc:
            st.error(T["no_desc"])
        else:
            with st.status(T["generating"], expanded=True) as status:
                try:
                    client = openai.OpenAI(api_key=API_KEY, base_url=BASE_URL)

                    if st.session_state.lang == "zh":
                        user_prompt = (
                            f"请根据以下信息生成 8D 报告："
                            f"产品：{product_name or '未提供'}, "
                            f"客户：{customer or '未提供'}, "
                            f"日期：{occur_date}, "
                            f"数量：{defect_qty}, "
                            f"严重程度：{severity}, "
                            f"标准：{industry_std}, "
                            f"团队：{team_members or '未提供'}\n\n"
                            f"问题描述：{problem_desc}"
                        )
                    else:
                        user_prompt = (
                            f"Generate 8D report based on:\n"
                            f"Product: {product_name or 'N/A'}\n"
                            f"Customer: {customer or 'N/A'}\n"
                            f"Date: {occur_date}\n"
                            f"Quantity: {defect_qty}\n"
                            f"Severity: {severity}\n"
                            f"Standard: {industry_std}\n"
                            f"Team: {team_members or 'N/A'}\n\n"
                            f"Problem Description: {problem_desc}"
                        )

                    response = client.chat.completions.create(
                        model="deepseek-chat",
                        messages=[
                            {"role": "system", "content": SYSTEM_PROMPT[st.session_state.lang]},
                            {"role": "user", "content": user_prompt}
                        ],
                        stream=True,
                        temperature=0.2,
                        max_tokens=4096
                    )

                    full_content = ""
                    stream_placeholder = st.empty()

                    for chunk in response:
                        delta = chunk.choices[0].delta.content
                        if delta:
                            full_content += delta
                            stream_placeholder.markdown(full_content)

                    status.update(label="✅ " + T["success"], state="complete", expanded=False)

                    final_result = clean_format(full_content)
                    st.session_state.current_result = final_result
                    inc_trial_used(user_id)
                    save_report_history(user_id, product_name, customer, problem_desc, final_result, st.session_state.lang)

                except openai.APIConnectionError:
                    status.update(label="❌ 网络连接失败", state="error")
                    st.error("🌐 网络连接失败，请检查网络后重试")
                except openai.RateLimitError:
                    status.update(label="❌ API 频率超限", state="error")
                    st.error("⏱️ API 调用频率超限，请等待 30 秒后重试")
                except openai.AuthenticationError:
                    status.update(label="❌ API 密钥验证失败", state="error")
                    st.error("🔑 API 密钥验证失败，请联系管理员")
                except openai.APIError as e:
                    status.update(label="❌ 服务异常", state="error")
                    st.error(f"❌ 服务异常：{getattr(e, 'type', '')}" if hasattr(e, 'type') else "❌ 服务异常，请稍后重试")
                except Exception as e:
                    status.update(label="❌ 系统错误", state="error")
                    logging.error(f"生成报告未知错误：{e}", exc_info=True)
                    st.error(T["api_error"])

with col_preview:
    st.header(T["preview_header"])
    if st.session_state.current_result:
        edit_mode = st.checkbox(T["edit_mode"], key="edit_mode_toggle")
        if edit_mode:
            edited = st.text_area(T["edit_placeholder"], value=st.session_state.current_result, height=600, key="edit_area")
            if st.button(T["save_edit"], use_container_width=True, key="save_edit_btn"):
                st.session_state.current_result = edited
                st.session_state.edit_mode_toggle = False
                st.success("✅ 修改已保存" if st.session_state.lang == "zh" else "✅ Changes saved")
                st.rerun()
        else:
            st.markdown(st.session_state.current_result.replace("**", "").replace("#", ""))
        
        st.markdown("---")
        user_id = st.session_state.get("user_id")
        lic = get_user_license(user_id) if user_id else None
        if lic and lic['plan_type'] != 'free':
            word_data = export_to_word(
                st.session_state.current_result,
                product_name or "8D_Report",
                st.session_state.lang
            )
            st.download_button(
                label=T["download_btn"],
                data=word_data,
                file_name=f"8D_Report_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        else:
            st.info(T["export_disabled"])
    else:
        st.info("👈 输入问题描述后点击生成" if st.session_state.lang == "zh" else "👈 Enter description and click generate")
